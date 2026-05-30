from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Pt


def compliance_dataframe(bundle: dict[str, Any]) -> pd.DataFrame:
    rows = []
    for item in bundle.get("compliance_matrix", []) or []:
        rows.append(
            {
                "Requirement ID": item.get("requirement_id", ""),
                "Requirement": item.get("requirement", ""),
                "Status": item.get("status", ""),
                "Confidence": item.get("confidence", ""),
                "Evidence": item.get("evidence", ""),
                "Response Strategy": item.get("response_strategy", ""),
                "Owner": item.get("owner", ""),
            }
        )
    return pd.DataFrame(rows)


def risk_dataframe(bundle: dict[str, Any]) -> pd.DataFrame:
    return pd.DataFrame(bundle.get("risks", []) or [])


def markdown_report(bundle: dict[str, Any]) -> str:
    lines = []
    lines.append(f"# {bundle.get('project_title', 'RFP Analysis')}\n")
    lines.append(f"**Bid recommendation:** {bundle.get('bid_recommendation', 'Needs More Info')}  ")
    lines.append(f"**Bid score:** {bundle.get('bid_score', 'N/A')}/100  ")
    lines.append(f"**Quality score:** {bundle.get('quality_score', 'N/A')}/100\n")

    lines.append("## Opportunity Summary\n")
    lines.append(str(bundle.get("opportunity_summary", "")) + "\n")

    lines.append("## Win Themes\n")
    for theme in bundle.get("win_themes", []) or []:
        lines.append(f"- {theme}")
    lines.append("")

    lines.append("## Requirement Matrix\n")
    for req in bundle.get("requirements", []) or []:
        lines.append(f"- **{req.get('id','')}** [{req.get('priority','')}/{req.get('category','')}]: {req.get('text','')}")
    lines.append("")

    lines.append("## Compliance Matrix\n")
    for item in bundle.get("compliance_matrix", []) or []:
        lines.append(f"### {item.get('requirement_id','')} — {item.get('status','Unknown')} ({item.get('confidence','')}%)")
        lines.append(f"**Requirement:** {item.get('requirement','')}")
        lines.append(f"**Evidence:** {item.get('evidence','')}")
        lines.append(f"**Strategy:** {item.get('response_strategy','')}\n")

    lines.append("## Risks\n")
    for risk in bundle.get("risks", []) or []:
        lines.append(f"- **{risk.get('severity','Medium')} / {risk.get('bid_impact','Minor')}**: {risk.get('risk','')} — {risk.get('mitigation','')}")
    lines.append("")

    lines.append("## Clarifying Questions\n")
    for q in bundle.get("clarifying_questions", []) or []:
        lines.append(f"- {q}")
    lines.append("")

    lines.append("## Solution Architecture\n")
    lines.append(str(bundle.get("solution_architecture", "")) + "\n")

    lines.append("## Delivery Plan\n")
    for m in bundle.get("delivery_plan", []) or []:
        lines.append(f"- **{m.get('milestone','')}** ({m.get('duration','')}): {m.get('outputs','')} Dependencies: {m.get('dependencies','')}")
    lines.append("")

    lines.append("## Pricing Assumptions\n")
    for a in bundle.get("pricing_assumptions", []) or []:
        lines.append(f"- {a}")
    lines.append("")

    lines.append("## Proposal Draft\n")
    lines.append(str(bundle.get("proposal_draft", "")) + "\n")

    lines.append("## Reviewer Notes\n")
    for note in bundle.get("reviewer_notes", []) or []:
        lines.append(f"- {note}")

    return "\n".join(lines).strip() + "\n"


def write_docx(bundle: dict[str, Any], path: str | Path) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    doc.add_heading(bundle.get("project_title", "RFP Analysis"), level=0)
    doc.add_paragraph(f"Bid recommendation: {bundle.get('bid_recommendation', 'Needs More Info')}")
    doc.add_paragraph(f"Bid score: {bundle.get('bid_score', 'N/A')}/100 | Quality score: {bundle.get('quality_score', 'N/A')}/100")

    def heading(text: str, level: int = 1):
        doc.add_heading(text, level=level)

    heading("Opportunity Summary")
    doc.add_paragraph(str(bundle.get("opportunity_summary", "")))

    heading("Win Themes")
    for theme in bundle.get("win_themes", []) or []:
        doc.add_paragraph(str(theme), style="List Bullet")

    heading("Compliance Matrix")
    df = compliance_dataframe(bundle)
    if not df.empty:
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr[i].text = str(col)
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(df.columns):
                cells[i].text = str(row[col])[:1200]
    else:
        doc.add_paragraph("No compliance items generated.")

    heading("Risks")
    for r in bundle.get("risks", []) or []:
        doc.add_paragraph(f"{r.get('severity','Medium')}: {r.get('risk','')}", style="List Bullet")
        doc.add_paragraph(f"Mitigation: {r.get('mitigation','')}")

    heading("Clarifying Questions")
    for q in bundle.get("clarifying_questions", []) or []:
        doc.add_paragraph(str(q), style="List Number")

    heading("Solution Architecture")
    doc.add_paragraph(str(bundle.get("solution_architecture", "")))

    heading("Delivery Plan")
    for m in bundle.get("delivery_plan", []) or []:
        doc.add_paragraph(f"{m.get('milestone','')} ({m.get('duration','')}): {m.get('outputs','')}", style="List Bullet")

    heading("Proposal Draft")
    for para in str(bundle.get("proposal_draft", "")).split("\n"):
        if para.startswith("# "):
            heading(para.replace("# ", ""), level=1)
        elif para.startswith("## "):
            heading(para.replace("## ", ""), level=2)
        elif para.startswith("- "):
            doc.add_paragraph(para[2:], style="List Bullet")
        elif para.strip():
            doc.add_paragraph(para)

    doc.save(path)
    return path


def save_all_outputs(bundle: dict[str, Any], out_dir: str | Path, slug: str) -> dict[str, Path]:
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    md = out_dir / f"{slug}_report.md"
    js = out_dir / f"{slug}_analysis.json"
    csv = out_dir / f"{slug}_compliance_matrix.csv"
    docx = out_dir / f"{slug}_proposal.docx"

    md.write_text(markdown_report(bundle), encoding="utf-8")
    js.write_text(json.dumps(bundle, indent=2, ensure_ascii=False), encoding="utf-8")
    compliance_dataframe(bundle).to_csv(csv, index=False)
    write_docx(bundle, docx)
    return {"markdown": md, "json": js, "csv": csv, "docx": docx}
