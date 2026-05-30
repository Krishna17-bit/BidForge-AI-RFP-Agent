from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import BinaryIO

from pypdf import PdfReader
from docx import Document as DocxDocument


@dataclass
class LoadedDocument:
    name: str
    text: str
    pages: list[dict]


def _read_txt(file_obj: BinaryIO) -> str:
    raw = file_obj.read()
    if isinstance(raw, str):
        return raw
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(enc)
        except Exception:
            continue
    return raw.decode("utf-8", errors="ignore")


def load_uploaded_file(uploaded_file) -> LoadedDocument:
    """Load PDF, DOCX or TXT/MD from a Streamlit UploadedFile."""
    name = uploaded_file.name
    suffix = Path(name).suffix.lower()
    uploaded_file.seek(0)

    if suffix == ".pdf":
        reader = PdfReader(uploaded_file)
        pages = []
        parts = []
        for i, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            pages.append({"page": i, "text": text})
            parts.append(f"\n\n--- {name} | page {i} ---\n{text}")
        return LoadedDocument(name=name, text="".join(parts).strip(), pages=pages)

    if suffix == ".docx":
        doc = DocxDocument(uploaded_file)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Include table text because RFPs often contain compliance tables.
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)
        text = "\n".join(paragraphs)
        return LoadedDocument(name=name, text=text, pages=[{"page": 1, "text": text}])

    if suffix in {".txt", ".md", ".csv"}:
        text = _read_txt(uploaded_file)
        return LoadedDocument(name=name, text=text, pages=[{"page": 1, "text": text}])

    raise ValueError(f"Unsupported file type: {suffix}. Please upload PDF, DOCX, TXT, MD or CSV.")


def load_path(path: str | Path) -> LoadedDocument:
    path = Path(path)
    with path.open("rb") as f:
        class FileWrapper:
            name = path.name
            def read(self, *args): return f.read(*args)
            def seek(self, *args): return f.seek(*args)
            def tell(self): return f.tell()
        return load_uploaded_file(FileWrapper())
